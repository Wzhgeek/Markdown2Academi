"""
Markdown 转 Word 转换器
基于 Pandoc 和 python-docx
"""

import os
import tempfile
import subprocess
import re
from typing import Optional, Dict, Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class MarkdownToDocxConverter:
    """Markdown 转 Word 转换器"""

    def __init__(self):
        self.templates = {
            "thesis": ThesisTemplate(),
            "journal": JournalTemplate(),
        }

    def convert(self, input_file: str, output_file: str, template: str = "thesis",
                metadata: Optional[Dict[str, Any]] = None):
        """
        转换 Markdown 文件到 Word

        Args:
            input_file: 输入 Markdown 文件路径
            output_file: 输出 Word 文件路径
            template: 模板名称 (thesis/journal)
            metadata: 额外的元数据
        """
        # 读取原始 Markdown
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # 预处理 Markdown 扩展语法
        md_content = self._preprocess_markdown(md_content)

        # 临时文件
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write(md_content)
            temp_md = f.name

        try:
            # 第一步：使用 pandoc 进行基础转换
            temp_docx = tempfile.mktemp(suffix='.docx')
            self._run_pandoc(temp_md, temp_docx)

            # 第二步：使用 python-docx 进行后处理
            doc = Document(temp_docx)

            # 应用模板样式
            template_handler = self.templates.get(template, ThesisTemplate())
            template_handler.apply(doc, self._extract_metadata(md_content))

            # 保存最终文档
            doc.save(output_file)

        finally:
            # 清理临时文件
            if os.path.exists(temp_md):
                os.unlink(temp_md)
            if os.path.exists(temp_docx):
                os.unlink(temp_docx)

    def _run_pandoc(self, input_file: str, output_file: str):
        """运行 pandoc 命令"""
        try:
            cmd = [
                'pandoc',
                input_file,
                '-o', output_file,
                '--from', 'markdown+yaml_metadata_block+citations',
                '--to', 'docx',
                '--reference-doc=default',
            ]
            subprocess.run(cmd, check=True, capture_output=True)
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"Pandoc 转换失败: {e.stderr.decode()}")
        except FileNotFoundError:
            raise RuntimeError("未找到 pandoc，请先安装: https://pandoc.org/installing.html")

    def _preprocess_markdown(self, content: str) -> str:
        """预处理 Markdown 扩展语法"""
        # 处理 #abstract
        content = re.sub(r'^#abstract\s*(.+?)(?=\n#|\Z)',
                         r'\n**摘要**\n\n\1\n', content, flags=re.MULTILINE | re.DOTALL)

        # 处理 #abstract-en
        content = re.sub(r'^#abstract-en\s*(.+?)(?=\n#|\Z)',
                         r'\n**Abstract**\n\n\1\n', content, flags=re.MULTILINE | re.DOTALL)

        # 处理 #keywords
        content = re.sub(r'^#keywords\s*(.+)$',
                         r'**关键词：** \1', content, flags=re.MULTILINE)

        # 处理 #figure
        content = re.sub(r'^#figure\s+(.+?)\s*\|\s*(.+?)\s*(?:\|\s*(.+))?$',
                         self._replace_figure, content, flags=re.MULTILINE)

        # 处理 #table
        content = re.sub(r'^#table\s+(.+?)\s*\|\s*(.+?)\s*(?:\|\s*(.+))?$',
                         self._replace_table, content, flags=re.MULTILINE)

        # 处理 #equation
        content = re.sub(r'^#equation\s+(.+?)(?:\s*\|\s*label=(.+))?$',
                         self._replace_equation, content, flags=re.MULTILINE)

        return content

    def _replace_figure(self, match) -> str:
        """替换图片标记"""
        caption = match.group(1)
        image_path = match.group(2)
        options = match.group(3) or ""

        width = "80%"
        if "width=" in options:
            width_match = re.search(r'width=(\d+%)', options)
            if width_match:
                width = width_match.group(1)

        return f'\n![{caption}]({image_path}){{ width={width} }}\n\n*{caption}*\n'

    def _replace_table(self, match) -> str:
        """替换表格标记"""
        caption = match.group(1)
        data_path = match.group(2)
        options = match.group(3) or ""

        # 如果是 CSV 文件，读取并转换为 Markdown 表格
        if data_path.endswith('.csv'):
            try:
                with open(data_path, 'r', encoding='utf-8') as f:
                    lines = f.readlines()

                header = "header=true" in options
                table_md = self._csv_to_markdown(lines, header)
                return f'\n{table_md}\n*{caption}*\n'
            except:
                pass

        return f'\n*{caption}*\n'

    def _csv_to_markdown(self, lines: list, has_header: bool = True) -> str:
        """CSV 转 Markdown 表格"""
        if not lines:
            return ""

        rows = [line.strip().split(',') for line in lines if line.strip()]
        if not rows:
            return ""

        md_lines = []
        # 表头
        md_lines.append('| ' + ' | '.join(rows[0]) + ' |')
        # 分隔符
        md_lines.append('| ' + ' | '.join(['---'] * len(rows[0])) + ' |')

        # 数据行
        start_idx = 1 if has_header else 0
        for row in rows[start_idx:]:
            md_lines.append('| ' + ' | '.join(row) + ' |')

        return '\n'.join(md_lines)

    def _replace_equation(self, match) -> str:
        """替换公式标记"""
        equation = match.group(1)
        label = match.group(2)

        if label:
            return f'\n$$ {equation} \\tag{{{label}}} $$\n'
        return f'\n$$ {equation} $$\n'

    def _extract_metadata(self, content: str) -> Dict[str, Any]:
        """提取 YAML 元数据"""
        metadata = {}
        yaml_match = re.match(r'^---\s*\n(.*?)\n---\s*\n', content, re.DOTALL)

        if yaml_match:
            yaml_content = yaml_match.group(1)
            for line in yaml_content.split('\n'):
                if ':' in line:
                    key, value = line.split(':', 1)
                    metadata[key.strip()] = value.strip().strip('"').strip("'")

        return metadata


class BaseTemplate:
    """基础模板类"""

    def apply(self, doc: Document, metadata: Dict[str, Any]):
        """应用模板样式"""
        raise NotImplementedError

    def set_heading_styles(self, doc: Document):
        """设置标题样式"""
        # 标题 1 - 章标题
        style = doc.styles['Heading 1']
        font = style.font
        font.name = '黑体'
        font.size = Pt(16)
        font.bold = True
        paragraph_format = style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format.space_before = Pt(24)
        paragraph_format.space_after = Pt(18)

        # 标题 2 - 节标题
        style = doc.styles['Heading 2']
        font = style.font
        font.name = '黑体'
        font.size = Pt(14)
        font.bold = True

        # 标题 3
        style = doc.styles['Heading 3']
        font = style.font
        font.name = '黑体'
        font.size = Pt(12)
        font.bold = True

    def set_normal_style(self, doc: Document):
        """设置正文样式"""
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.first_line_indent = Inches(0.5)


class ThesisTemplate(BaseTemplate):
    """毕业论文模板"""

    def apply(self, doc: Document, metadata: Dict[str, Any]):
        """应用毕业论文样式"""
        self.set_heading_styles(doc)
        self.set_normal_style(doc)

        # 添加封面（如果文档为空）
        if len(doc.paragraphs) == 0 or not doc.paragraphs[0].text.strip():
            self._add_cover_page(doc, metadata)

        # 添加页眉页脚
        self._add_headers_footers(doc, metadata)

        # 处理特殊段落（摘要、关键词等）
        self._process_special_paragraphs(doc)

    def _add_cover_page(self, doc: Document, metadata: Dict[str, Any]):
        """添加封面"""
        title = metadata.get('title', '论文标题')
        author = metadata.get('author', '作者')
        school = metadata.get('school', '学院')

        # 学校名称
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(school)
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.name = '黑体'

        # 论文类型
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('本科毕业论文')
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.name = '黑体'

        doc.add_paragraph()  # 空行

        # 论文标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = '宋体'

        doc.add_paragraph()
        doc.add_paragraph()

        # 作者信息表格样式
        info_items = [
            ('学　　院', school),
            ('专　　业', metadata.get('major', '')),
            ('学生姓名', author),
            ('学　　号', metadata.get('student_id', '')),
            ('指导教师', metadata.get('advisor', '')),
        ]

        for label, value in info_items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'{label}：{value}')
            run.font.size = Pt(14)
            run.font.name = '宋体'

        doc.add_page_break()

    def _add_headers_footers(self, doc: Document, metadata: Dict[str, Any]):
        """添加页眉页脚"""
        for section in doc.sections:
            # 页眉
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = metadata.get('school', '')
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.runs[0].font.size = Pt(10)

            # 页脚 - 页码
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加页码字段
            footer_para.add_run()._element.append(
                self._create_page_number_field(footer_para)
            )

    def _create_page_number_field(self, paragraph):
        """创建页码字段"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        paragraph._p.append(fldChar1)
        paragraph._p.append(instrText)
        paragraph._p.append(fldChar2)
        paragraph._p.append(fldChar3)

        return fldChar3

    def _process_special_paragraphs(self, doc: Document):
        """处理特殊段落样式"""
        for para in doc.paragraphs:
            text = para.text.strip()

            # 摘要标题
            if text == '摘要' or text == 'Abstract':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(16)
                    run.font.bold = True
                    run.font.name = '黑体'

            # 关键词
            elif text.startswith('关键词：') or text.startswith('Keywords:'):
                for run in para.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
                    if '关键词' in text or 'Keywords' in text:
                        run.font.bold = True


class JournalTemplate(BaseTemplate):
    """期刊论文模板"""

    def apply(self, doc: Document, metadata: Dict[str, Any]):
        """应用期刊论文样式"""
        self.set_heading_styles(doc)
        self.set_normal_style(doc)

        # 期刊通常不需要复杂页眉页脚
        # 但可能需要特定的引用格式

        # 设置双倍行距（某些期刊要求）
        for para in doc.paragraphs:
            para.paragraph_format.line_spacing = 2.0

        # 处理标题
        title = metadata.get('title', '')
        if title and doc.paragraphs:
            first_para = doc.paragraphs[0]
            if title in first_para.text:
                first_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in first_para.runs:
                    run.font.size = Pt(14)
                    run.font.bold = True
